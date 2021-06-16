from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

def replace_text(changes_list, template_doc_name, new_doc_name):
    document = Document(template_doc_name)

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    for item in changes_list:
        print(item[0])
        for paragraph in document.paragraphs:
            if item[0] in paragraph.text:
                paragraph.text = paragraph.text.replace(item[0], str(item[1]))
                paragraph.style = document.styles['Normal']
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if item[0] in paragraph.text:
                            paragraph.text = paragraph.text.replace(item[0], str(item[1]))
                            paragraph.style = document.styles['Normal']
        
        # отдельно переделывается нижний колонтитул
        style = document.styles['Footer']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(8)

        section = document.sections[0]
        for paragraph in section.footer.paragraphs:
            if item[0] in paragraph.text:
                paragraph.text = paragraph.text.replace(item[0], str(item[1]))
                paragraph.style = document.styles['Footer']

        new_name = new_doc_name + '.docx'

    document.save(new_name)
    return new_name